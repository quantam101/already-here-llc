#!/usr/bin/env python3
"""Generate a combined Word proposal from the RFP response."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

out = Path(__file__).parent / "PROPOSAL.docx"
doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


add_heading("Arizona Multi-Site Electronics Asset Inventory & 2D QR Tagging", level=1)
add_paragraph("Proposal prepared for Jeff Pomeroy, Arizona Multi-Site (AZ-MVS)", bold=True)
add_paragraph("Prepared by: Already Here LLC")
add_paragraph("Date: August 11, 2026")
add_paragraph("Reply to: dispatch@alreadyherellc.com")
add_paragraph("Phone: (602) 882-2920")
add_paragraph("Web: www.alreadyherellc.com")
doc.add_paragraph()

add_heading("Cover Letter", level=2)
cover = (
    "Dear Jeff,\n\n"
    "Thank you for inviting Already Here LLC to bid on the Arizona multi-site electronics asset inventory and 2D QR tagging project. We reviewed the RFP and the master location schedule, and we are ready to execute this turnkey baseline across all 24 physical site visits.\n\n"
    "Our proposal includes a wall-to-wall physical inventory of qualifying electronics at each site, durable scannable 2D QR asset tags with unique human-readable asset IDs, data capture and nightly reconciliation against the RFP required data fields, and a consolidated enterprise master asset register in Excel and CSV with location-level files, an exception report, site completion certifications, and a final project summary.\n\n"
    "Our base one-time price for the full scope is $13,600 (assuming approximately 700 tagged assets; the site-visit and processing components are fixed, and tags are priced per actual unit). This is an all-in cost of approximately $19.43 per tagged asset at the 700-asset estimate. Pricing is anchored to our published rate card and compared against local small-business IT service rates, not national firms.\n\n"
    "We can mobilize within one week of a signed SOW and confirmed site schedule. I am available for a call to walk through the methodology, staffing model, and tag sample.\n\n"
    "Best regards,\n\n"
    "Stephen Franklin\n"
    "Owner / Field Operations\n"
    "Already Here LLC\n"
    "dispatch@alreadyherellc.com\n"
    "(602) 882-2920"
)
for line in cover.split("\n"):
    if line.strip() == "":
        doc.add_paragraph()
    else:
        add_paragraph(line)

doc.add_page_break()

add_heading("1. Project Overview", level=2)
add_paragraph(
    "Already Here LLC will deliver a turnkey wall-to-wall physical inventory of qualifying electronics and technology assets across 24 physical site visits in Arizona, apply durable unique 2D QR code asset tags, capture the data fields required by the RFP, perform quality assurance and reconciliation, and deliver a consolidated enterprise master asset register with location-level files, an exception log, site completion certifications, and a final project summary."
)
add_paragraph(
    "The engagement is scoped from kickoff through final deliverables, with all work performed during normal business hours and coordinated to minimize disruption to customer-facing operations."
)

add_heading("2. Scope of Work", level=2)
scope_items = [
    "Pre-deployment planning: confirm site schedule, field procedures, tag format, asset inclusion rules, data template, site contacts, security/access requirements, and escalation process.",
    "Onsite physical inventory: room-by-room and area-by-area sweep at each location, including offices, counters, storage rooms, network/IT closets, back-of-house areas, manager offices, and other areas containing qualifying assets.",
    "2D QR asset tagging: supply and apply durable, tamper-resistant 2D QR labels with unique human-readable asset IDs. Tags placed consistently without covering manufacturer labels, ventilation, service panels, or operational controls.",
    "Asset data capture: capture all fields listed in the RFP, including Asset ID, QR code, business/entity, physical location, address, area/room, category, manufacturer, model, serial/service tag, MAC address (where accessible), description, condition, assigned user, photo reference, status, inventory date, and notes/exceptions.",
    "Asset photographs: one clear photograph per qualifying asset over $500 replacement value or as designated by Client; filename keyed to Asset ID.",
    "Exception identification: flag assets that are damaged, obsolete, unused, stored, missing identifying information, inaccessible, duplicated, or otherwise require follow-up.",
    "Quality assurance: reconcile tag counts to inventory records, review duplicate/missing serials and asset IDs, perform a spot-check of not less than 10% of tagged assets per site, and provide a site completion/exception report before closing each location.",
    "Final deliverables: consolidated enterprise master asset register plus location-level inventory files in Excel and CSV, data dictionary, exception log, site completion certifications, and project completion summary.",
    "PawnCo P4 option: if exercised, a separate pre-opening sweep after equipment is staged/installed, tagged and registered at $500 per sweep plus $0.75 per tag.",
]
for item in scope_items:
    doc.add_paragraph(item, style="List Bullet")

add_paragraph("Out of scope: low-value consumables and accessories unless specifically directed; disconnecting, moving, or altering production equipment without express authorization; after-hours work without written approval; travel beyond the Phoenix metro area without a written travel addendum.")

add_heading("3. Asset Tag Specification", level=2)
tag_items = [
    "Material: white polyester thermal-transfer labels with permanent acrylic adhesive.",
    "Size: 1\" x 2\" (25 mm x 51 mm).",
    "Print: black 2D QR code plus human-readable Asset ID.",
    "Service life: 5+ years under normal indoor commercial conditions.",
    "Tamper-evident: destructible or void-pattern option available for high-touch assets.",
    "QR content: unique internal asset identifier only (no URLs, network credentials, or confidential data). Optional hosted lookup available under the optional software platform.",
    "Replacement: lost or damaged tags replaced at $0.75 each during the project; after project closeout, replacement tags available at the same rate plus a dispatch fee.",
]
for item in tag_items:
    doc.add_paragraph(item, style="List Bullet")

add_heading("4. Methodology, Staffing, and Quality Assurance", level=2)
method_steps = [
    "Kickoff call: review scope, site list, access requirements, asset inclusion rules, and data template.",
    "Tag sample: produce and ship a sample tag for Client approval before field production.",
    "Field execution: one or two trained Already Here technicians per site, using mobile scanning/capture devices. Data is synced to a secure project workspace nightly.",
    "QA/reconciliation: back-office review for duplicate Asset IDs, missing serials, model mismatches, and photo linkage; on-site spot-check before closeout.",
    "Closeout: site completion certification signed by the field lead; exception log reviewed with Client contact.",
    "Final delivery: master register, location-level files, exception report, photo package, and project summary delivered via secure file transfer.",
]
for step in method_steps:
    doc.add_paragraph(step, style="List Number")

add_paragraph(
    "Estimated duration: 4-5 weeks from kickoff to final deliverables. Estimated field days: 12-16 business days (2-3 sites per day). Estimated assets per technician per day: 30-50 assets, varying by site density and access."
)

add_heading("5. Site Schedule", level=2)
add_paragraph("The following 24 physical site visits are based on the RFP master location schedule.")

sites = [
    ["MVS 05", "Motor Vehicle", "940 N Alma School Rd, Suite 105", "Chandler", "85224", "Operating"],
    ["MVS 10", "Motor Vehicle", "1501 N 16th St", "Phoenix", "85006", "Operating"],
    ["MVS 12", "Motor Vehicle", "833 N Cooper Rd, Suite 101", "Gilbert", "85233", "Operating"],
    ["MVS 13", "Motor Vehicle", "11249 W Buckeye Rd", "Avondale", "85323", "Operating"],
    ["MVS 14", "Motor Vehicle", "4385 W Bell Rd", "Glendale", "85308", "Operating"],
    ["MVS 16", "Motor Vehicle", "3655 W Anthem Way, Suite B115", "Anthem", "85086", "Operating"],
    ["MVS 20", "Motor Vehicle", "2525 N Scottsdale Rd, Suite 1-3", "Scottsdale", "85257", "Operating"],
    ["MVS 22", "Motor Vehicle", "21805 S Ellsworth Rd, Suite B107", "Queen Creek", "85142", "Operating"],
    ["MVS 23", "Motor Vehicle + Insurance Satellite", "17019 W Greenway Rd, Suite 114", "Surprise", "85388", "Insurance embedded"],
    ["MVS 24", "Motor Vehicle", "29455 N Cave Creek Rd, Suite 126", "Cave Creek", "85331", "Operating"],
    ["MVS 26", "Motor Vehicle", "20783 N 83rd Ave, Suite 105", "Peoria", "85382", "Operating"],
    ["MVS 27", "Motor Vehicle", "6740 E University Dr, Suite 106", "Mesa", "85205", "Operating"],
    ["MVS 28", "Motor Vehicle", "10720 W Indian School Rd, Suite 51", "Phoenix", "85037", "Operating"],
    ["MVS 29", "Motor Vehicle", "2620 W Baseline Rd", "Mesa", "85202", "Operating"],
    ["MVS 32", "Motor Vehicle", "3172 E Indian School Rd", "Phoenix", "85016", "Operating"],
    ["MVS 33", "Motor Vehicle", "85 W Combs Rd, Suite 103", "San Tan Valley", "85140", "Operating"],
    ["MVS 34", "Motor Vehicle", "1010 W Chandler Heights Rd", "Chandler", "85248", "Operating"],
    ["MVS 51 / Motor Carrier", "Two MVS Ops + Insurance Satellite", "5036 W Cactus Rd, Suites 3 & 4", "Glendale", "85304", "MVS 51 = Suite 3; Motor Carrier = Suite 4; insurance embedded"],
    ["PawnCo P1", "Pawn Shop", "40975 N Ironwood Dr, Unit B111", "San Tan Valley", "85140", "Operating"],
    ["PawnCo P2", "Pawn Shop", "13220 W Van Buren St, Suite 100", "Goodyear", "85338", "Operating"],
    ["PawnCo P3", "Pawn Shop", "21041 N Cave Creek Rd", "Phoenix", "85024", "Operating"],
    ["PawnCo P4 - Thunderbird", "Pawn Shop", "5140 W Thunderbird Rd", "Glendale", "-", "Under Construction"],
    ["The Grove Restaurant & Bar", "Restaurant", "12555 W Bell Rd", "Surprise", "85378", "Operating"],
    ["Corporate Office", "Corporate", "2633 E Indian School Rd", "Phoenix", "85016", "Operating"],
]

headers = ["Site / Operation", "Type", "Address", "City", "ZIP", "Status / Notes"]
table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for p in hdr_cells[i].paragraphs:
        for run in p.runs:
            run.font.bold = True
for row in sites:
    cells = table.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v

doc.add_paragraph()
add_heading("6. Pricing", level=2)
pricing_rows = [
    ["Project management / kickoff", "1 project", "—", "$1,500.00"],
    ["Onsite inventory & tagging", "24 site visits", "$450.00/visit", "$10,800.00"],
    ["Asset tags", "Per tag (actual)", "$0.75/tag", "~$525.00 (est. 700 tags)"],
    ["Data capture & reconciliation", "Included + final QA pass", "$750.00 fixed", "$750.00"],
    ["Photos", "Included per asset", "Included", "Included"],
    ["Travel / mobilization", "Phoenix metro", "Included", "Included"],
    ["PawnCo P4 pre-opening option", "1 site option", "$500 + tags", "$500.00 (if exercised)"],
    ["TOTAL ONE-TIME PROJECT PRICE (base)", "—", "—", "$13,575.00"],
]

ptable = doc.add_table(rows=1, cols=4)
ptable.style = "Table Grid"
phdr = ptable.rows[0].cells
for i, h in enumerate(["Pricing Component", "Qty / Basis", "Unit Price", "Extended / Fixed Price"]):
    phdr[i].text = h
    for p in phdr[i].paragraphs:
        for run in p.runs:
            run.font.bold = True
for row in pricing_rows:
    cells = ptable.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v

doc.add_paragraph()
add_paragraph(
    "Assumed total asset count: approximately 700 tagged assets across 24 physical site visits. Estimated all-in cost per tagged asset: ~$19.43 at 700 assets; ~$17.00 at 800 assets; ~$22.67 at 600 assets. Rounded proposal total: $13,600. Proposal validity: 30 days. Payment terms: Net 15, 50% on kickoff and 50% on final delivery. Base scope not to exceed $14,750 without a written change order."
)

add_heading("7. Optional Services", level=2)
opt_rows = [
    ["Hosted asset-management platform", "$1,200/year or $100/month", "Web register, mobile QR scanning, CSV/XLSX export, non-proprietary data export guaranteed."],
    ["Annual physical re-inventory", "$7,200/year ($300/site x 24)", "Validates asset moves, updates register, replaces missing/damaged tags."],
    ["PawnCo P4 pre-opening sweep", "$500 per visit + $0.75/tag", "Performed after equipment is staged/installed."],
]
otable = doc.add_table(rows=1, cols=3)
otable.style = "Table Grid"
ohdr = otable.rows[0].cells
for i, h in enumerate(["Optional Service", "Pricing", "Notes"]):
    ohdr[i].text = h
    for p in ohdr[i].paragraphs:
        for run in p.runs:
            run.font.bold = True
for row in opt_rows:
    cells = otable.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v

add_heading("8. Deliverables", level=2)
deliverables = [
    "Pre-deployment implementation plan and confirmed site schedule.",
    "Finalized asset-tag specification and a sample tag for Client approval before production tagging.",
    "Complete enterprise master asset register in .xlsx and .csv format.",
    "Separate location-level asset registers (or filterable Site/Location field in the master file).",
    "Asset photographs, where included, and an unambiguous method for linking photos to Asset IDs.",
    "Exception report identifying inaccessible, untagged, damaged, obsolete, duplicate, or unresolved assets.",
    "Site completion certification for each location showing assets verified, assets tagged, and exceptions.",
    "Final project summary with total assets inventoried/tagged by site and category.",
]
for item in deliverables:
    doc.add_paragraph(item, style="List Bullet")

add_heading("9. Assumptions", level=2)
assumptions = [
    "Safe site access, badging, escort (if required), and a point of contact are available during the scheduled window.",
    "Work is performed during normal business hours (Monday-Friday, 08:00-17:00 MST) unless otherwise agreed.",
    "Client confirms asset inclusion rules and any capitalization/value threshold at kickoff.",
    "Qualifying assets are powered on or physically accessible without disassembly; labels that cannot be read without disruption are recorded as exceptions.",
    "Physical site count remains at 24 visits; additional visits require a change order.",
    "Travel is within the Phoenix metropolitan area; sites more than 60 miles one-way from downtown Phoenix may require a travel addendum.",
    "Final asset count is expected to be in the 600-800 range; per-tag pricing reconciles to actual count.",
]
for item in assumptions:
    doc.add_paragraph(item, style="List Bullet")

add_heading("10. Acceptance", level=2)
add_paragraph(
    "Deliverables are accepted when Client signs the completion certificate or does not reject the final register in writing within five (5) business days of delivery. Minor corrections (formatting, missing fields, or exception clarifications) will be resolved within two (2) business days at no additional charge."
)

add_heading("11. Insurance and Compliance", level=2)
add_paragraph(
    "Already Here LLC maintains commercial general liability insurance of at least $1,000,000 per occurrence and will provide a certificate of insurance on request. Provider personnel will follow Client site security, safety, confidentiality, and data-handling policies made available in advance."
)

add_heading("12. Signatures", level=2)
sig_table = doc.add_table(rows=3, cols=2)
sig_table.style = "Table Grid"
sig_table.rows[0].cells[0].text = "Already Here LLC"
sig_table.rows[0].cells[1].text = "Arizona Multi-Site (AZ-MVS)"
sig_table.rows[1].cells[0].text = "Signature: ______________________"
sig_table.rows[1].cells[1].text = "Signature: ______________________"
sig_table.rows[2].cells[0].text = "Name / Date: Stephen Franklin, Owner / ____________"
sig_table.rows[2].cells[1].text = "Name / Date: Jeff Pomeroy / ____________"

add_heading("13. About Already Here LLC", level=2)
add_paragraph(
    "Already Here LLC is a Phoenix-anchored field operations company providing smart hands, network, POS, data center, retail technology, asset-tagging, and infrastructure support across Arizona. We work with MSPs, vendors, prime contractors, and multi-site operators who need reliable onsite execution with structured closeout documentation."
)
add_paragraph("Representative multi-site engagements (client names available under NDA on request):")
refs = [
    "National QSR chain — POS hardware installation across 4 Mesa, AZ metro locations with photo-documented closeout.",
    "National retail brand — RFID reader survey (55 readers, 4 APs, 61 data runs) with full structured field execution and documentation.",
    "Enterprise infrastructure vendor — HPE Alletra MP deployment in a Chandler, AZ data center, multi-day engagement with rack, cable, and verification.",
    "Retail technology program — Store technology recovery and closeout support, including device identification, inventory capture, packing, labeling, and return-logistics handoff.",
]
for ref in refs:
    doc.add_paragraph(ref, style="List Bullet")

doc.save(str(out))
print(f"Generated {out}")
